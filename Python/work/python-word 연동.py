# 코드 3-11 워드 문서 생성 및 파일로 저장

# python-docx 패키지 설치
from docx import Document
filename = '워드문서1.docx'
document = Document()
document.save(filename)

# 코드 3-12 워드 문서 생성 및 문단 추가
from docx import Document
filename = '워드문서2.docx'
document = Document()
document.add_paragraph('첫 번째 문단 입니다.')
document.add_paragraph('두 번째 문단 입니다.')
document.add_paragraph('세 번째 문단 입니다.')
document.save(filename)


# 코드 3-13 첫 번째 문단 출력
from docx import Document
filename = '워드문서2.docx'
# document 객체 생성
document = Document(filename)
# 문단 목록
paragraphs = document.paragraphs
# 첫번째 문단
first_paragraph = paragraphs[0]
# 첫번째 문단 내 문자열 출력
print(first_paragraph.text)


# 코드 3-14 문단에 스타일 적용
from docx import Document
# font 관련 내용을 다루는 Pt 클래스 import
from docx.shared import Pt
from docx.oxml.ns import qn
filename = '워드문서2.docx'
# document 객체 생성
document = Document(filename)
# 첫번째 문단의 첫 번째 런에 접근 (문단 전체가 하나의 런)
first_paragraph = document.paragraphs[0]
first_run = first_paragraph.runs[0]
# 해당 런에 다양한 스타일 적용
first_run.italic = True
first_run.underline = True
first_run.bold = True
first_run.font.size = Pt(24)
# 한글 폰트
first_run.font.name = 'HY견고딕'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY견고딕')
# strike: True인 경우 문자 중앙에 라인이 그러짐.
first_run.font.strike = True
filename = '워드문서3.docx'
document.save(filename)


# 코드 3-15 문단에 새로운 스타일의 텍스트 추가
from docx import Document
# font 관련 내용을 다루는 Pt 클래스 import
from docx.shared import Pt
from docx.oxml.ns import qn
filename = '워드문서3.docx'
# document 객체 생성
document = Document(filename)
# 두번째 문단의 첫 번째 런에 접근 (문단 전체가 하나의 런)
second_paragraph = document.paragraphs[1]
new_run = second_paragraph.add_run('새로운 스타일의 텍스트')
# 해당 런에 다양한 스타일 적용
new_run.italic = True
new_run.underline = True
new_run.bold = True
new_run.font.size = Pt(18)
# 한글 폰트
new_run.font.name = 'HY목각파임B'
new_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'HY목각파임B')
filename = '워드문서4.docx'
document.save(filename)


# 코드 3-16 문단에 그림 추가
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
filename = '워드문서4.docx'
# document 객체 생성
document = Document(filename)
# 새로운 문단 추가
new_paragraph = document.add_paragraph()
# 새로운 문단의 가운데 정렬
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 새로운 문단에 새로운 런을 추가
logo_run = new_paragraph.add_run()
# 새로 추가한 런에 그림 추가
logo_run.add_picture('plot.png')
# 그림을 추가한 다음 줄 바꾸기를 적용
logo_run.add_break(WD_BREAK.LINE)
# 그림 설명을 위한 런 추가
caption_run = new_paragraph.add_run('[그림 비에이퍼블릭 로고]')
filename = '워드문서5.docx'
document.save(filename)


# 코드 3-17 문단에 테이블 추가
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
filename = '워드문서5.docx'
# document 객체 생성
document = Document(filename)
# 2행 2열 크기의 테이블 추가
table = document.add_table(rows=2, cols=2)
# 테이블의 테두리 그려줌. default는 테두리 없음.
table.style = 'Table Grid'
# 1행 1열의 셀 접근
table.cell(0, 0).text = '1행 1열'
# 1행 2열의 셀 접근과 문단 가져오기
# paragraph = table.cell(0,1).paragraph[0]
paragraph = table.cell(0, 1).add_paragraph()
new_run = paragraph.add_run('1행 2열')
new_run.bold = True
# 텍스트 색상 변경
new_run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
# 2행 1열의 셀 접근 과 새로운 문단 추가
new_paragraph = table.cell(1, 0).add_paragraph()
new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_run = new_paragraph.add_run('2행 1열')
new_run.italic = True
new_run.font.color.rgb = RGBColor(0, 0, 255)
# 2행 2열의 셀 접근 과 새로운 문단 추가
# paragraph = table.cell(1,1).add_paragraph[0]
paragraph = table.cell(1, 1).add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
new_run = paragraph.add_run('2행 2열 첫 번째')
new_run.italic = True
new_run.font.color.rgb = RGBColor(0x00, 0xff, 0x00)
new_paragraph = table.cell(1, 1).add_paragraph()
new_run = new_paragraph.add_run('2행 2열 두 번째')
new_run.bold = True
new_run.italic = True
new_run.font.color.rgb = RGBColor(0x66, 0x00, 0xcc)
filename = '워드문서6.docx'
document.save(filename)


# ch3-19 규칙파일 내용 읽어 오기
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx import Document
def get_template(filename):
    file = open(filename, 'r')
    logo = file.readline().split(':')[1].strip()
    title = file.readline().split(':')[1].strip()
    content = file.readline().split(':')[1].strip()
    agendas = file.readline().split(':')[1].strip()
    date = file.readline().split(':')[1].strip()
    author = file.readline().split(':')[1].strip()

    return logo, title, content, agendas, date, author


logo, title, content, agendas, date, author = get_template('report2.txt')
logo
title
content
agendas
date
author


# python-docx 설치
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx import Document

def get_template(filename):
    file = open(filename, 'r')
    logo = file.readline().split(':')[1].strip()
    title = file.readline().split(':')[1].strip()
    content = file.readline().split(':')[1].strip()
    agendas = file.readline().split(':')[1].strip()
    date = file.readline().split(':')[1].strip()
    author = file.readline().split(':')[1].strip()

    return logo, title, content, agendas, date, author

# 로고 그림을 추가하는 함수
def add_logo(document, logo_path):
    logo_p = document.add_paragraph()
    # logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    logo_run = logo_p.add_run()
    logo_run.add_picture(logo_path, width=Cm(3), height=Cm(3))
    logo_run.add_break(WD_BREAK.LINE)
    logo_run.add_break(WD_BREAK.LINE)

# 타이틀을 추가하는 함수
def add_title(document, title):
    title_p = document.add_paragraph()
    # title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    # title_run.font_name = '새굴림'
    title_run.font.name = '맑은 고딕'
    # title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    title_run.add_break(WD_BREAK.LINE)

# 본문을 추가하는 함수
def add_content(document, content):
    content_p = document.add_paragraph()
    # content_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    content_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    content_run = content_p.add_run(content)
    content_run.font.size = Pt(12)
    # content_run.font_name = '새굴림'
    content_run.font.name = '맑은 고딕'
    # content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    content_run.add_break(WD_BREAK.LINE)

# 안건을 추가하는 함수
def add_agendas(document, agendas):
    for agenda in agendas:
        agenda_p = document.add_paragraph(agenda, style='List Bullet')
        agenda_run = agenda_p.runs[0]
        agenda_run.font.size = Pt(12)
        # agenda_run.font_name = '새굴림'
        agenda_run.font.name = '맑은 고딕'
        # agenda_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
        agenda_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 날짜를 추가하는 함수
def add_date(document, date):
    date_p = document.add_paragraph()
    # date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(date)
    date_run.font.size = Pt(12)
    # date_run.font_name = '새굴림'
    date_run.font.name = '맑은 고딕'
    # date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    date_run.add_break(WD_BREAK.LINE)

logo, title, content, agendas, date, author = get_template('report2.txt')

filename = 'report2.docx'
document = Document()

add_logo(document, logo)
add_title(document, title)
add_content(document, content)
add_agendas(document, agendas.split(','))
add_date(document, date)
add_date(document, author)

filename = 'report3.docx'
document.save(filename)

# python-docx 설치

# from docx.shared import Cm
# from docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.text import WD_BREAK
# from docx import Document

from docx.shared import Cm
from docx.shared import Pt
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.enum.text import WD_BREAK
from docx import Document


def get_template(filename):
    file = open(filename, 'r')
    logo = file.readline().split(':')[1].strip()

    title = file.readline().split(':')[1].strip()
    content = file.readline().split(':')[1].strip()
    agendas = file.readline().split(':')[1].strip()
    date = file.readline().split(':')[1].strip()
    author = file.readline().split(':')[1].strip()

    return logo, title, content, agendas, date, author


# 로고 그림을 추가하는 함수
def add_logo(document, logo_path):
    logo_p = document.add_paragraph()
    # logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    logo_run = logo_p.add_run()
    logo_run.add_picture(logo_path, width=Cm(3), height=Cm(3))
    logo_run.add_break(WD_BREAK.LINE)
    logo_run.add_break(WD_BREAK.LINE)


# 타이틀을 추가하는 함수
def add_title(document, title):
    title_p = document.add_paragraph()
    # title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    # title_run.font_name = '새굴림'
    title_run.font.name = '맑은 고딕'
    # title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    title_run.add_break(WD_BREAK.LINE)


# 본문을 추가하는 함수
def add_content(document, content):
    content_p = document.add_paragraph()
    # content_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    content_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    content_run = content_p.add_run(content)
    content_run.font.size = Pt(12)
    # content_run.font_name = '새굴림'
    content_run.font.name = '맑은 고딕'
    # content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    content_run.add_break(WD_BREAK.LINE)


# 안건을 추가하는 함수
def add_agendas(document, agendas):
    for agenda in agendas:
        agenda_p = document.add_paragraph(agenda, style='List Bullet')
        agenda_run = agenda_p.runs[0]
        agenda_run.font.size = Pt(12)
        # agenda_run.font_name = '새굴림'
        agenda_run.font.name = '맑은 고딕'
        # agenda_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
        agenda_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


# 날짜를 추가하는 함수
def add_date(document, date):
    date_p = document.add_paragraph()
    # date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(date)
    date_run.font.size = Pt(12)
    # date_run.font_name = '새굴림'
    date_run.font.name = '맑은 고딕'
    # date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '새굴림')
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    date_run.add_break(WD_BREAK.LINE)


def main():
    logo, title, content, agendas, date, author = get_template('report2.txt')

    filename = 'report2.docx'
    document = Document()

    add_logo(document, logo)
    add_title(document, title)
    add_content(document, content)
    add_agendas(document, agendas.split(','))
    add_date(document, date)
    add_date(document, author)

    document.save(filename)

# if __name__ == "__main__":
#     main()

main()



