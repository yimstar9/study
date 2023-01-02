#임성구
#아래 이미지에 보이는 것처럼 다음 조건에 맞게 워드파일 내에
#텍스트와 이미지 파일을 추가하는 python 코드를 작성하시오.
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
import os
os.chdir('E:\\GoogleDrive\\2022년 빅데이터&AI 강의\\workplace\\Python\\work')
os.getcwd()

filename = 'YimSungGoo.docx'
document = Document()


# 1) 첫 번째 문단
# “파이썬을 활용한 Word문서 작성” 텍스트
# Font: 맑은 고딕
# Font size: 24
# Bold

#워드 문서 생성 및 문단 추가
document.add_paragraph('파이썬을 활용한 Word문서 작성')
document.add_paragraph('임성구')

# 첫번째 문단의 첫 번째 런에 접근 (문단 전체가 하나의 런)
first_paragraph = document.paragraphs[0]
first_run = first_paragraph.runs[0]

# 해당 런에 다양한 스타일 적용
first_run.bold = True
first_run.font.size = Pt(24)

# 한글 폰트
first_run.font.name = '맑은고딕'
first_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은고딕')


# 2) 두 번째 문단
# “본인이름” 텍스트
# Font: 굴림
# Font size: 20
# Underline

# 두번째 문단의 첫 번째 런에 접근 (문단 전체가 하나의 런)
second_paragraph = document.paragraphs[1]
second_run = second_paragraph.runs[0]

# 해당 런에 다양한 스타일 적용
second_run.underline = True
second_run.font.size = Pt(20)

# 한글 폰트
second_run.font.name = '굴림'
second_run._element.rPr.rFonts.set(qn('w:eastAsia'), '굴림')


# 3) 세 번째 문단
# Tjoeun_logo2.jpg 이미지 파일 첨부
# 중앙 정렬
# 그림 설명 부분 [더조은 컴퓨터 아카데미] 추가

# 세번째 문단 추가
third_paragraph = document.add_paragraph()

# 세번째 문단의 가운데 정렬
third_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 새로운 문단에 새로운 런을 추가
logo_run = third_paragraph.add_run()

# 새로 추가한 런에 그림 추가
logo_run.add_picture('tjoeun_logo2.jpg')

# 그림을 추가한 다음 줄 바꾸기를 적용
logo_run.add_break(WD_BREAK.LINE)

# 그림 설명을 위한 런 추가
caption_run = third_paragraph.add_run('[더조은 컴퓨터 아카데미]')


# 4) 네 번째 문단
# “훈련중입니다.” 텍스트
# Font: 맑은고딕
# Font size: 40
# 이탤릭체

document.add_paragraph('훈련중입니다.')

# 네번째 문단의 첫 번째 런에 접근
fourth_paragraph = document.paragraphs[3]
fourth_run = fourth_paragraph.runs[0]

# 해당 런에 다양한 스타일 적용
fourth_run.italic = True
fourth_run.font.size = Pt(40)

# 한글 폰트
fourth_run.font.name = '맑은고딕'
fourth_run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은고딕')

document.save(filename)